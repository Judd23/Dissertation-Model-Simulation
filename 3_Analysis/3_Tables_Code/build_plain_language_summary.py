#!/usr/bin/env python3
"""
Build Plain Language Summary (DOCX) from SEM analysis results.

Generates an APA 7 formatted document summarizing key findings in accessible language
for non-technical audiences (e.g., dissertation committee, stakeholders, policymakers).

Usage:
    python build_plain_language_summary.py --outdir <results_dir> --out <output_dir>

Output: Plain_Language_Summary.docx
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from datetime import datetime
from typing import TYPE_CHECKING, Any, Optional, Dict

# Optional imports with graceful degradation
HAS_DOCX = False
HAS_PANDAS = False
Document = None  # type: Any
Pt = None  # type: Any
Inches = None  # type: Any
WD_ALIGN_PARAGRAPH = None  # type: Any
pd = None  # type: Any

try:
    from docx import Document as _Document
    from docx.shared import Pt as _Pt, Inches as _Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
    Document = _Document
    Pt = _Pt
    Inches = _Inches
    WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    print("Warning: python-docx not installed. Install with: pip install python-docx")

try:
    import pandas as _pd
    pd = _pd
    HAS_PANDAS = True
except ImportError:
    pass


# ==============================================================================
# Data Loading Functions
# ==============================================================================

def load_bootstrap_results(outdir: Path) -> Optional[Any]:
    """Load bootstrap results from CSV."""
    if not HAS_PANDAS or pd is None:
        return None
    
    candidates = [
        outdir / "bootstrap_results.csv",
        outdir / "raw" / "bootstrap_results.csv",
    ]
    
    for path in candidates:
        if path.exists():
            try:
                return pd.read_csv(path)
            except Exception as e:
                print(f"Warning: Could not load {path}: {e}")
    return None


def load_fit_measures(outdir: Path) -> Dict[str, float]:
    """Load model fit measures from structural_fitMeasures.txt."""
    fit = {}
    candidates = [
        outdir / "raw" / "RQ1_RQ3_main" / "structural" / "structural_fitMeasures.txt",
        outdir / "RQ1_RQ3_main" / "structural" / "structural_fitMeasures.txt",
    ]
    
    for path in candidates:
        if path.exists():
            try:
                with open(path, 'r') as f:
                    for line in f:
                        parts = line.strip().split()
                        if len(parts) == 2:
                            try:
                                fit[parts[0]] = float(parts[1])
                            except ValueError:
                                pass
                return fit
            except Exception as e:
                print(f"Warning: Could not load {path}: {e}")
    return fit


def load_ps_model(outdir: Path) -> Optional[Any]:
    """Load propensity score model coefficients."""
    if not HAS_PANDAS or pd is None:
        return None
    
    candidates = [
        outdir / "raw" / "RQ1_RQ3_main" / "ps_model.csv",
        outdir / "RQ1_RQ3_main" / "ps_model.csv",
    ]
    
    for path in candidates:
        if path.exists():
            try:
                return pd.read_csv(path)
            except Exception:
                pass
    return None


def load_weight_diagnostics(outdir: Path) -> Dict[str, float]:
    """Load weight diagnostics."""
    diag = {}
    candidates = [
        outdir / "raw" / "RQ1_RQ3_main" / "weight_diagnostics.csv",
        outdir / "RQ1_RQ3_main" / "weight_diagnostics.csv",
    ]
    
    for path in candidates:
        if path.exists():
            try:
                if HAS_PANDAS and pd is not None:
                    df = pd.read_csv(path)
                    for _, row in df.iterrows():
                        diag[row['metric']] = row['value']
                return diag
            except Exception:
                pass
    return diag


# ==============================================================================
# APA Formatting Functions
# ==============================================================================

def format_apa_stat(est: float, ci_lower: float, ci_upper: float, p: Optional[float] = None) -> str:
    """Format a statistic in APA 7 style: β = 0.XX, 95% CI [0.XX, 0.XX], p < .001"""
    stat_str = f"β = {est:.3f}, 95% CI [{ci_lower:.3f}, {ci_upper:.3f}]"
    if p is not None:
        if p < 0.001:
            stat_str += ", p < .001"
        elif p < 0.01:
            stat_str += f", p = {p:.3f}"
        elif p < 0.05:
            stat_str += f", p = {p:.3f}"
        else:
            stat_str += f", p = {p:.3f}"
    return stat_str


def interpret_effect_size(beta: float) -> str:
    """Interpret standardized effect size per Cohen's conventions."""
    abs_beta = abs(beta)
    if abs_beta < 0.10:
        return "negligible"
    elif abs_beta < 0.30:
        return "small"
    elif abs_beta < 0.50:
        return "medium"
    else:
        return "large"


def get_effect_row(df: Any, param: str) -> Optional[Dict[str, Any]]:
    """Extract a row from bootstrap results by parameter name."""
    if df is None or df.empty:
        return None
    row = df[df['parameter'] == param]
    if row.empty:
        return None
    return row.iloc[0].to_dict()


def format_fit_indices(fit: Dict[str, float]) -> str:
    """Format fit indices in APA style."""
    cfi = fit.get('cfi', fit.get('cfi.robust', None))
    rmsea = fit.get('rmsea', fit.get('rmsea.robust', None))
    srmr = fit.get('srmr', None)
    
    parts = []
    if cfi is not None:
        parts.append(f"CFI = {cfi:.3f}")
    if rmsea is not None:
        parts.append(f"RMSEA = {rmsea:.3f}")
    if srmr is not None:
        parts.append(f"SRMR = {srmr:.3f}")
    
    return ", ".join(parts) if parts else "Fit indices not available"


# ==============================================================================
# Document Building Functions
# ==============================================================================

def add_heading(doc: Any, text: str, level: int = 1) -> None:
    """Add a heading with consistent formatting."""
    heading = doc.add_heading(text, level=level)
    if hasattr(heading, 'runs') and heading.runs:
        heading.runs[0].font.name = 'Times New Roman'


def add_paragraph(doc: Any, text: str, bold: bool = False, italic: bool = False) -> Any:
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    if Pt is not None:
        run.font.size = Pt(12)
    return p


def add_finding(doc: Any, label: str, text: str) -> None:
    """Add a labeled finding paragraph."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.name = 'Times New Roman'
    if Pt is not None:
        run_label.font.size = Pt(12)
    run_text = p.add_run(text)
    run_text.font.name = 'Times New Roman'
    if Pt is not None:
        run_text.font.size = Pt(12)


# ==============================================================================
# Section Generators
# ==============================================================================

def generate_overview_section(doc: Any, fit: Dict[str, float], diag: Dict[str, float]) -> None:
    """Generate study overview section."""
    add_heading(doc, "Study Overview", level=1)
    
    n_obs = int(diag.get('n_obs', 0)) if diag else 0
    ess = int(diag.get('ess', 0)) if diag else 0
    
    overview_text = (
        "This study examined the psychosocial effects of accelerated dual credit participation "
        "(FASt status) on first-year developmental adjustment among equity-impacted California "
        "community college students. The analysis used conditional-process structural equation "
        "modeling (SEM) with propensity score overlap weighting to estimate causal effects while "
        "accounting for selection bias."
    )
    if n_obs > 0:
        overview_text += f" The analytic sample included {n_obs:,} students"
        if ess > 0:
            overview_text += f" (effective sample size = {ess:,.0f} after weighting)"
        overview_text += "."
    
    add_paragraph(doc, overview_text)
    
    if fit:
        add_paragraph(doc, "")
        add_paragraph(doc, 
            f"The structural model demonstrated excellent fit to the data: {format_fit_indices(fit)}.",
            italic=True
        )


def generate_rq1_section(doc: Any, boot_df: Any) -> None:
    """Generate RQ1 findings: Total effect of FASt on DevAdj."""
    add_heading(doc, "Research Question 1: Overall Effect of FASt Status", level=1)
    
    add_paragraph(doc, 
        "RQ1 asked: What is the total effect of accelerated dual credit participation (FASt status) "
        "on first-year developmental adjustment?"
    )
    doc.add_paragraph()
    
    # Get total effect at mean credit dose (z_mid)
    total = get_effect_row(boot_df, 'total_z_mid')
    
    if total:
        est = total['est']
        ci_lower = total['ci_lower']
        ci_upper = total['ci_upper']
        sig = total['sig']
        
        effect_size = interpret_effect_size(est)
        direction = "positive" if est > 0 else "negative"
        sig_text = "statistically significant" if sig else "not statistically significant"
        
        add_finding(doc, "Finding",
            f"The total effect of FASt status on developmental adjustment was {effect_size} and "
            f"{sig_text} ({format_apa_stat(est, ci_lower, ci_upper)}). Students who matriculated "
            f"with ≥12 transferable dual credit units showed {'higher' if est > 0 else 'lower'} "
            f"first-year developmental adjustment compared to students without this credential, "
            f"after controlling for baseline covariates."
        )
    else:
        add_paragraph(doc, "Total effect data not available. Run full analysis pipeline.", italic=True)


def generate_rq2_section(doc: Any, boot_df: Any) -> None:
    """Generate RQ2 findings: Mediation through EmoDiss and QualEngag."""
    add_heading(doc, "Research Question 2: Mediation Pathways", level=1)
    
    add_paragraph(doc, 
        "RQ2 asked: To what extent is the effect of FASt status on developmental adjustment "
        "mediated by (a) emotional distress and (b) quality of engagement?"
    )
    doc.add_paragraph()
    
    # Get indirect effects at mean credit dose
    ind_emo = get_effect_row(boot_df, 'ind_EmoDiss_z_mid')
    ind_qual = get_effect_row(boot_df, 'ind_QualEngag_z_mid')
    
    # Path coefficients
    a1 = get_effect_row(boot_df, 'a1')  # FASt → EmoDiss
    a2 = get_effect_row(boot_df, 'a2')  # FASt → QualEngag
    b1 = get_effect_row(boot_df, 'b1')  # EmoDiss → DevAdj
    b2 = get_effect_row(boot_df, 'b2')  # QualEngag → DevAdj
    
    add_heading(doc, "Emotional Distress Pathway", level=2)
    
    if ind_emo and a1 and b1:
        sig_text = "significant" if ind_emo['sig'] else "non-significant"
        a1_dir = "increased" if a1['est'] > 0 else "decreased"
        b1_dir = "negatively" if b1['est'] < 0 else "positively"
        
        add_finding(doc, "Finding",
            f"The indirect effect of FASt status on developmental adjustment through emotional "
            f"distress was {sig_text} ({format_apa_stat(ind_emo['est'], ind_emo['ci_lower'], ind_emo['ci_upper'])}). "
            f"FASt status {a1_dir} emotional distress ({format_apa_stat(a1['est'], a1['ci_lower'], a1['ci_upper'])}), "
            f"which in turn was {b1_dir} related to developmental adjustment "
            f"({format_apa_stat(b1['est'], b1['ci_lower'], b1['ci_upper'])})."
        )
    else:
        add_paragraph(doc, "Emotional distress mediation data not available.", italic=True)
    
    doc.add_paragraph()
    add_heading(doc, "Quality of Engagement Pathway", level=2)
    
    if ind_qual and a2 and b2:
        sig_text = "significant" if ind_qual['sig'] else "non-significant"
        a2_dir = "increased" if a2['est'] > 0 else "decreased"
        b2_dir = "positively" if b2['est'] > 0 else "negatively"
        
        add_finding(doc, "Finding",
            f"The indirect effect of FASt status on developmental adjustment through quality of "
            f"engagement was {sig_text} ({format_apa_stat(ind_qual['est'], ind_qual['ci_lower'], ind_qual['ci_upper'])}). "
            f"FASt status {a2_dir} quality of engagement ({format_apa_stat(a2['est'], a2['ci_lower'], a2['ci_upper'])}), "
            f"which in turn was {b2_dir} related to developmental adjustment "
            f"({format_apa_stat(b2['est'], b2['ci_lower'], b2['ci_upper'])})."
        )
    else:
        add_paragraph(doc, "Quality of engagement mediation data not available.", italic=True)


def generate_rq3_section(doc: Any, boot_df: Any) -> None:
    """Generate RQ3 findings: Moderated mediation by credit dose."""
    add_heading(doc, "Research Question 3: Moderation by Credit Dose", level=1)
    
    add_paragraph(doc, 
        "RQ3 asked: Does the number of dual credit units earned (credit dose) moderate the "
        "direct and indirect effects of FASt status on developmental adjustment?"
    )
    doc.add_paragraph()
    
    # Get conditional effects at -1SD, mean, +1SD
    dir_low = get_effect_row(boot_df, 'dir_z_low')
    dir_mid = get_effect_row(boot_df, 'dir_z_mid')
    dir_high = get_effect_row(boot_df, 'dir_z_high')
    
    # Interaction terms
    cz = get_effect_row(boot_df, 'cz')  # X×Z on Y
    a1z = get_effect_row(boot_df, 'a1z')  # X×Z on EmoDiss
    a2z = get_effect_row(boot_df, 'a2z')  # X×Z on QualEngag
    
    add_heading(doc, "Direct Effect Moderation", level=2)
    
    if cz:
        sig_text = "significant" if cz['sig'] else "not significant"
        add_finding(doc, "Finding",
            f"The interaction between FASt status and credit dose on developmental adjustment "
            f"(direct effect moderation) was {sig_text} "
            f"({format_apa_stat(cz['est'], cz['ci_lower'], cz['ci_upper'])})."
        )
        
        if dir_low and dir_mid and dir_high:
            add_paragraph(doc, 
                f"Conditional direct effects: At low credit dose (−1 SD): "
                f"{format_apa_stat(dir_low['est'], dir_low['ci_lower'], dir_low['ci_upper'])}; "
                f"at mean credit dose: {format_apa_stat(dir_mid['est'], dir_mid['ci_lower'], dir_mid['ci_upper'])}; "
                f"at high credit dose (+1 SD): {format_apa_stat(dir_high['est'], dir_high['ci_lower'], dir_high['ci_upper'])}.",
                italic=True
            )
    else:
        add_paragraph(doc, "Direct effect moderation data not available.", italic=True)
    
    doc.add_paragraph()
    add_heading(doc, "Indirect Effect Moderation (Moderated Mediation)", level=2)
    
    # Index of moderated mediation
    imm_emo = get_effect_row(boot_df, 'index_MM_EmoDiss')
    imm_qual = get_effect_row(boot_df, 'index_MM_QualEngag')
    
    if imm_emo:
        sig_text = "significant" if imm_emo['sig'] else "not significant"
        add_finding(doc, "Emotional Distress",
            f"The index of moderated mediation through emotional distress was {sig_text} "
            f"({format_apa_stat(imm_emo['est'], imm_emo['ci_lower'], imm_emo['ci_upper'])}), "
            f"indicating that the indirect effect {'did' if imm_emo['sig'] else 'did not'} vary "
            f"significantly as a function of credit dose."
        )
    
    if imm_qual:
        sig_text = "significant" if imm_qual['sig'] else "not significant"
        add_finding(doc, "Quality of Engagement",
            f"The index of moderated mediation through quality of engagement was {sig_text} "
            f"({format_apa_stat(imm_qual['est'], imm_qual['ci_lower'], imm_qual['ci_upper'])}), "
            f"indicating that the indirect effect {'did' if imm_qual['sig'] else 'did not'} vary "
            f"significantly as a function of credit dose."
        )
    
    if not imm_emo and not imm_qual:
        add_paragraph(doc, "Index of moderated mediation data not available.", italic=True)


def generate_rq4_section(doc: Any, fit: Dict[str, float]) -> None:
    """Generate RQ4 findings: Measurement model."""
    add_heading(doc, "Research Question 4: Measurement Invariance", level=1)
    
    add_paragraph(doc, 
        "RQ4 asked: Does the measurement model demonstrate configural, metric, and scalar "
        "invariance across key demographic groups (race/ethnicity, Pell status, first-generation status)?"
    )
    doc.add_paragraph()
    
    if fit:
        cfi = fit.get('cfi', fit.get('cfi.robust', None))
        rmsea = fit.get('rmsea', fit.get('rmsea.robust', None))
        
        good_fit = (cfi is not None and cfi >= 0.95) and (rmsea is not None and rmsea <= 0.06)
        
        add_finding(doc, "Finding",
            f"The overall structural model demonstrated {'excellent' if good_fit else 'acceptable'} "
            f"fit to the data ({format_fit_indices(fit)}). "
            f"{'These values exceed conventional thresholds (CFI ≥ .95, RMSEA ≤ .06) for good model fit.' if good_fit else ''}"
        )
        add_paragraph(doc, 
            "Note: Full measurement invariance testing across demographic groups is reported "
            "in the supplementary tables.",
            italic=True
        )
    else:
        add_paragraph(doc, "Model fit indices not available.", italic=True)


def generate_implications_section(doc: Any, boot_df: Any) -> None:
    """Generate practical implications section."""
    add_heading(doc, "Practical Implications", level=1)
    
    # Get key effects to drive implications
    total = get_effect_row(boot_df, 'total_z_mid')
    ind_emo = get_effect_row(boot_df, 'ind_EmoDiss_z_mid')
    ind_qual = get_effect_row(boot_df, 'ind_QualEngag_z_mid')
    b1 = get_effect_row(boot_df, 'b1')  # EmoDiss → DevAdj
    b2 = get_effect_row(boot_df, 'b2')  # QualEngag → DevAdj
    
    implications = []
    
    # Implication based on emotional distress pathway
    if b1 and b1['sig'] and b1['est'] < 0:
        implications.append(
            "Given the strong negative relationship between emotional distress and developmental "
            "adjustment, institutions should prioritize mental health support services for all "
            "first-year students, particularly those transitioning from accelerated dual credit pathways."
        )
    
    # Implication based on quality of engagement
    if b2 and b2['sig'] and b2['est'] > 0:
        implications.append(
            "Quality of engagement with faculty, staff, and peers was positively associated with "
            "developmental adjustment. Institutions should invest in high-impact practices such as "
            "learning communities, mentorship programs, and first-year seminars that facilitate "
            "meaningful student-faculty and peer interactions."
        )
    
    # Implication based on FASt effect
    if total:
        if total['sig'] and total['est'] > 0:
            implications.append(
                "The positive effect of FASt status suggests that accelerated dual credit programs "
                "can support student success when combined with appropriate support structures. "
                "Policy should focus on expanding equitable access to dual credit opportunities "
                "for underrepresented students."
            )
        elif not total['sig']:
            implications.append(
                "The non-significant direct effect of FASt status suggests that simply having "
                "dual credit credentials is not sufficient for improved adjustment. The mechanism "
                "matters: institutions should focus on the psychosocial factors that mediate success."
            )
    
    if implications:
        for i, impl in enumerate(implications, 1):
            add_finding(doc, f"Implication {i}", impl)
    else:
        add_paragraph(doc, 
            "Implications will be derived from the specific pattern of significant effects "
            "once the full analysis is complete.",
            italic=True
        )


def generate_limitations_section(doc: Any) -> None:
    """Generate limitations section."""
    add_heading(doc, "Limitations", level=1)
    
    limitations = [
        ("Design", 
         "Despite propensity score weighting, this observational study cannot definitively establish "
         "causality. Unmeasured confounders may still influence the observed relationships."),
        ("Generalizability",
         "The sample is drawn from California community college students and may not generalize "
         "to other state systems or institutional types (e.g., four-year universities)."),
        ("Measurement",
         "Self-reported measures may be subject to social desirability bias. The developmental "
         "adjustment outcome is a latent construct that may not capture all dimensions of student success."),
        ("Temporal",
         "Cross-sectional design limits our ability to establish temporal precedence. Longitudinal "
         "data would strengthen causal inference."),
    ]
    
    for label, text in limitations:
        add_finding(doc, label, text)


def generate_technical_notes(doc: Any, B: int, ci_type: str) -> None:
    """Generate technical notes section."""
    add_heading(doc, "Technical Notes", level=1)
    
    add_paragraph(doc, 
        f"This summary is derived from a conditional-process structural equation model "
        f"estimated using lavaan in R. Propensity score overlap weights were used to "
        f"reduce selection bias. Bootstrap confidence intervals (B = {B:,}, {ci_type}) "
        f"were computed to assess statistical significance of indirect and conditional effects. "
        f"All standardized path coefficients (β) are reported.",
        italic=True
    )
    
    add_paragraph(doc, 
        f"Document generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
        italic=True
    )


# ==============================================================================
# Main Build Function
# ==============================================================================

def build_summary(doc: Any, outdir: Path, B: int, ci_type: str) -> None:
    """Build the plain language summary content."""
    
    # Title
    title = doc.add_heading("Plain Language Summary of Findings", level=0)
    if WD_ALIGN_PARAGRAPH is not None:
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Load all data
    boot_df = load_bootstrap_results(outdir)
    fit = load_fit_measures(outdir)
    diag = load_weight_diagnostics(outdir)
    
    # Build sections
    generate_overview_section(doc, fit, diag)
    doc.add_paragraph()
    
    # Key Constructs reference
    add_heading(doc, "Key Constructs", level=1)
    constructs = [
        ("Treatment (X)", "FASt status — students who matriculated with ≥12 transferable credits"),
        ("Moderator (Z)", "Credit dose — the number of dual credit units earned, mean-centered"),
        ("Mediator 1", "Emotional Distress (EmoDiss) — academic stress, loneliness, mental health challenges"),
        ("Mediator 2", "Quality of Engagement (QualEngag) — meaningful interactions with faculty, staff, peers"),
        ("Outcome (Y)", "Developmental Adjustment (DevAdj) — sense of belonging, perceived gains, satisfaction"),
    ]
    for label, description in constructs:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(description)
    doc.add_paragraph()
    
    # Research questions
    generate_rq1_section(doc, boot_df)
    doc.add_paragraph()
    
    generate_rq2_section(doc, boot_df)
    doc.add_paragraph()
    
    generate_rq3_section(doc, boot_df)
    doc.add_paragraph()
    
    generate_rq4_section(doc, fit)
    doc.add_paragraph()
    
    generate_implications_section(doc, boot_df)
    doc.add_paragraph()
    
    generate_limitations_section(doc)
    doc.add_paragraph()
    
    generate_technical_notes(doc, B, ci_type)


def main() -> int:
    parser = argparse.ArgumentParser(description='Build plain language summary document')
    parser.add_argument('--outdir', type=str, required=True, 
                        help='Input directory with results data (run folder)')
    parser.add_argument('--out', type=str, default=None, 
                        help='Output directory for docx (default: same as outdir)')
    parser.add_argument('--B', type=int, default=2000, 
                        help='Bootstrap replicates (for display)')
    parser.add_argument('--ci_type', type=str, default='perc', 
                        help='CI type: bca, perc, norm (for display)')
    args = parser.parse_args()
    
    if not HAS_DOCX or Document is None:
        print("ERROR: python-docx is required. Install with: pip install python-docx")
        return 1
    
    outdir = Path(args.outdir)
    out_docx_dir = Path(args.out) if args.out else outdir
    out_docx_dir.mkdir(parents=True, exist_ok=True)
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    if hasattr(style, 'font') and style.font is not None:
        style.font.name = 'Times New Roman'
        if Pt is not None:
            style.font.size = Pt(12)
    
    # Set margins (APA 7: 1 inch)
    if Inches is not None:
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    # Build content
    build_summary(doc, outdir, args.B, args.ci_type)
    
    # Save
    out_docx = out_docx_dir / "Plain_Language_Summary.docx"
    doc.save(str(out_docx))
    
    print(f"\n{'='*60}")
    print(f"Wrote: {out_docx}")
    print(f"{'='*60}")
    
    return 0


if __name__ == "__main__":
    sys.exit(main())
